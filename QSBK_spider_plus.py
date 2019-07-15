# coding=utf-8
#导入需要使用的模块
#import itertools  
import urllib
#import sys,io
#sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')#为了防止有时候输出产生中文乱码的问题
import requests
from lxml import etree
import json
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt
import docx.image
import os,shutil
class QSBK_aoto_spider():
    def __init__(self):        
        #url 和 User-Agent
        self.url = "https://www.qiushibaike.com/imgrank/page/{}/"#爬取含有图片的糗事百科
        self.headers = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/75.0.3770.100 Safari/537.36"}
    def get_url_list(self):
        url_list=[self.url.format(i) for i in range(1,3)]
        return url_list
    def parse_url(self,url):        
        print("正在处理：",url)
        response = requests.get(url,headers=self.headers)
        return response.content.decode()
    def get_info(self,html_str):
        html = etree.HTML(html_str)
        div_list = html.xpath("//div[@id='content-left']/div")
        content_list = []
        piccc = [] #初始化图片列表
        for div in div_list:
            info = {}
            picture = {}
            #info["用户"] = div.xpath(".//h2/text()")[0].strip() if len(div.xpath(".//h2/text()"))>0 else None
            info["内容"] = div.xpath(".//div[@class='content']/span/text()")
            info["内容"] = [i.strip().replace(" ", "") for i in info["内容"]]
            #info["好笑"] = div.xpath(".//span[@class='stats-vote']/i/text()")
            #info["好笑"] = info["好笑"][0] if len(info["好笑"])>0 else None
            #info["评论"] = div.xpath(".//span[@class='stats-comments']//i/text()")
            #info["评论"] =  info["评论"][0] if len(info["评论"])>0 else None
            picture["图片地址"] = div.xpath(".//div[@class='thumb']//img/@src")
            picture["图片地址"] = "https:"+picture["图片地址"][0] if len(picture["图片地址"])>0 else None
            #将字典内容附加到content_list
            #print(info["图片地址"])
            content_list.append(info)
            #将图片地址附加到piccc            
            piccc.append(picture["图片地址"])
        return piccc,content_list
    def save_info(self,content_list,piccc):# 保存数据
        global count_1,count_2,count_3
        for picsss in piccc:
            urllib.request.urlretrieve(picsss,Path+'/pictures/'+str(count_1)+'.jpg')
            count_1+=1
        for content in content_list:
            doc.add_paragraph(json.dumps(str(count_2)))
            doc.add_paragraph(json.dumps(content['内容'],ensure_ascii=False))
            doc.add_paragraph("**********************************************")
            doc.add_paragraph("**********************************************\n")
            try:
                doc.add_picture(Path+'/pictures/'+str(count_3)+'.jpg',width = Inches(6.0))
                doc.add_paragraph("\n")
                doc.add_paragraph("**********************************************")
                doc.add_paragraph("**********************************************\n")
            #处理图片有时候出现的异常 
            except docx.image.exceptions.UnrecognizedImageError:
                doc.add_picture('D:/code/zomkey.jpg',width = Inches(1.25))
                doc.add_paragraph('图片好像不见了哦~')                 
            doc.save(Path+'/contents/QSBK_aoto_spider.docx')
            count_2+=1
            count_3+=1
        print("保存成功")         
    def run(self):
        #1.根据url地址的规律,构造url list
        url_list = self.get_url_list()
        #2.发送请求,获取响应
        for url in url_list:
            html_str = self.parse_url(url)
            #3.提取数据
            piccc,content_list = self.get_info(html_str)            
            #4.保存
            self.save_info(content_list,piccc)
if __name__ == '__main__':
    count_1=1
    count_2=1
    count_3=1
    #判断是否已存在目录Path，如果存在则删除
    Path='D:/code/QSBK_auto_spider_plus'
    if os.path.exists(Path):
        shutil.rmtree(Path)
    #创建所需目录
    os.makedirs(Path)
    os.mkdir(Path+'/pictures')
    os.mkdir(Path+'/contents')
    #docx文件初始化
    doc = Document()
    doc.styles['Normal'].font.size=Pt(18)
    doc.styles['Normal'].font.name = u'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.add_heading(u"糗事百科热门精选-by Zomkey's Python_Spider",0)
    # 实例化方法
    QSBK = QSBK_aoto_spider()
    QSBK.run()


